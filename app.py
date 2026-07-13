"""
霖信莯咨询 · Qiaoxi Contract-Analyzer · 商业决策辅助系统
"""
import os, json, time, tempfile, traceback, threading, logging, re as _re
from datetime import datetime
import io as _io
import requests as _requests
from docx import Document as _Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
from openai import OpenAI

from src.fsm import FSM, State
from src.config import DEEPSEEK_API_KEY, DEEPSEEK_BASE_URL, DEEPSEEK_MODEL
from src.state1_parse import QiaoxiContractParser
from src.state2_legal_review import QiaoxiLegalReviewer
from src.state3_cld import CLDBuilder
from src.state4_council import CouncilRunner
from src.state5_simulation import SimulationEngine, DebateSynthesizer
from src.state6_decision import DecisionEngine
from src.state8_reconstruct import ReconstructionEngine
from src.security import (
    AuditLogger, deep_sanitize_contract_text, USER_CONSENT_TEMPLATE, sanitize_pii,
)

st.set_page_config(page_title="霖信莯 · 商业合同审查", page_icon="⚖️", layout="wide", initial_sidebar_state="collapsed")
llm = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)

st.markdown("""<style>
    .main-header { background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%); color: #fff; padding: 28px 32px; border-radius: 12px; margin-bottom: 24px; text-align: center; }
    .main-header h1 { font-size: 24px; font-weight: 700; letter-spacing: 2px; margin: 0; color: #fff; }
    .main-header p { font-size: 13px; opacity: .65; margin: 6px 0 0; color: #fff; }
    .risk-high { background: #fff1f0; border-left: 4px solid #cf1322; padding: 12px; margin: 8px 0; border-radius: 4px; }
    .risk-medium { background: #fff7e6; border-left: 4px solid #d46b08; padding: 12px; margin: 8px 0; border-radius: 4px; }
    .risk-low { background: #f6ffed; border-left: 4px solid #52c41a; padding: 12px; margin: 8px 0; border-radius: 4px; }
    .footer { position: fixed; bottom: 0; left: 0; right: 0; background: #1a1a2e; color: #aaa; text-align: center; padding: 10px; font-size: 12px; z-index: 1000; }
    .q-card { background: #f7f9fc; border-radius: 10px; padding: 20px; margin: 12px 0; border: 1px solid #e8ecf1; }
    .consent-box { background: #fafafa; border: 2px solid #0f3460; border-radius: 10px; padding: 24px; margin: 16px 0; max-height: 420px; overflow-y: auto; font-size: 13px; }
    .delete-countdown { background: #cf1322; color: #fff; padding: 12px 20px; border-radius: 8px; text-align: center; font-size: 18px; font-weight: 700; }
    textarea { display: none !important; }
    /* 解除复制限制 */
    * { -webkit-user-select: text !important; user-select: text !important; }
    /* 移除 Streamlit 工具栏干扰 */
    [data-testid="stElementToolbar"] { display: none !important; }
    [data-testid="stDecoration"] { display: none !important; }
    #MainMenu { display: none !important; }
    header[data-testid="stHeader"] { display: none !important; }
    [data-testid="stToolbar"] { display: none !important; }
</style>
<script>
(function() {
    document.addEventListener('copy', function(e) { e.stopImmediatePropagation(); }, true);
    document.addEventListener('selectstart', function(e) { e.stopImmediatePropagation(); }, true);
    document.addEventListener('contextmenu', function(e) { e.stopImmediatePropagation(); }, true);
})();
</script>
<div class="main-header">
    <h1>⚖️ 霖信莯 · 商业合同审查系统</h1>
    <p>霖信莯咨询 ｜ Qiaoxi Contract-Analyzer ｜ 商业决策辅助系统</p>
</div>
""", unsafe_allow_html=True)

# ─── Session State ───
DEFAULTS = {
    "phase": "consent",
    "base_paid": True,           # 收费机制已屏蔽（会员授权码由协会网站统一管控），保持机制代码不动
    "uploaded_file": None, "uploaded_file_path": None,
    "contract_summary": None, "clause_tree": None,
    "contract_raw": None,
    "contract_sanitized": None,
    "profile": None, "legal_review": None, "final_report": None,
    "r1_questions": None, "r1_answers": None,
    "r2_questions": None, "r2_answers": None,
    "error_msg": None, "consent_granted": False,
    "delete_requested": False,
    "session_files": [],
    # Phase 2 中间态
    "cld_report": None,
    "six_audits": None,
    "veto_any": False,
    "veto_auditors": [],
    "simulation": None,
    "debate": None,
    "decision": None,
    # 高级服务 / Phase 3
    "reconstruct_requested": False,
    "reconstruct_angle": None,
    "reconstruct_outputs": None,
    "reconstruct_paid": True,   # 收费机制已屏蔽，重构功能直接可用
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

audit = AuditLogger("audit.log")
logger = logging.getLogger("qiaoxi.app")

# ─── 底部免责声明 ───
st.markdown("""<div class="footer">
    ⚠️ 本系统为商业决策 <b>辅助</b> 工具，不构成正式法律意见或投资建议。最终决策由用户自行做出。<br/>
    昆明霖信莯科技有限公司 ｜ 合同文件存储于本地服务器 ｜ 霖信莯 · Qiaoxi Contract-Analyzer 仅接收脱敏后的信息<br/>
    <span style="color:#888;font-size:11px;">系统开发人员：李屹泉（身份证号：530111200801227358，联系电话：18987688373）</span>
</div>""", unsafe_allow_html=True)


# ─── 工具函数 ───

def _markdown_to_docx(md_text: str) -> bytes:
    """将 Markdown 文本转换为 docx 字节流，保留标题/加粗/引用/表格/列表结构。"""
    doc = _Document()

    # 全局字体设置
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    # 标题样式字号
    heading_sizes = {1: 18, 2: 15, 3: 13, 4: 12}

    def _add_heading(text: str, level: int):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.name = "微软雅黑"
        p.runs[0].font.size = Pt(heading_sizes.get(level, 12))

    def _parse_inline(para, text: str):
        """处理行内 **加粗** 标记，写入 para"""
        parts = _re.split(r'\*\*(.*?)\*\*', text)
        for j, part in enumerate(parts):
            run = para.add_run(part)
            run.font.name = "微软雅黑"
            run.font.size = Pt(11)
            if j % 2 == 1:  # 奇数片段是加粗内容
                run.bold = True

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 标题
        if line.startswith("#### "):
            _add_heading(line[5:].strip(), 4); i += 1; continue
        if line.startswith("### "):
            _add_heading(line[4:].strip(), 3); i += 1; continue
        if line.startswith("## "):
            _add_heading(line[3:].strip(), 2); i += 1; continue
        if line.startswith("# "):
            _add_heading(line[2:].strip(), 1); i += 1; continue

        # 水平分割线
        if line.strip() in ("---", "___", "***"):
            doc.add_paragraph("─" * 40); i += 1; continue

        # 引用块 >
        if line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            _parse_inline(p, line[2:].strip())
            i += 1; continue

        # 表格：连续的 | 行
        if line.startswith("|") and "|" in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                if not _re.match(r'^\|[-| :]+\|$', lines[i]):  # 跳过分隔行
                    table_lines.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            if table_lines:
                col_count = max(len(r) for r in table_lines)
                table = doc.add_table(rows=len(table_lines), cols=col_count)
                table.style = "Table Grid"
                for ri, row in enumerate(table_lines):
                    for ci, cell_text in enumerate(row):
                        cell = table.cell(ri, ci)
                        cell.text = cell_text
                        for run in cell.paragraphs[0].runs:
                            run.font.name = "微软雅黑"
                            run.font.size = Pt(10)
                            if ri == 0:
                                run.bold = True
            continue

        # 无序列表
        if _re.match(r'^[-*] ', line):
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline(p, line[2:].strip())
            i += 1; continue

        # 有序列表
        if _re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style="List Number")
            _parse_inline(p, _re.sub(r'^\d+\. ', '', line).strip())
            i += 1; continue

        # 空行
        if line.strip() == "":
            i += 1; continue

        # 普通段落
        p = doc.add_paragraph()
        _parse_inline(p, line.strip())
        i += 1

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def _schedule_file_deletion(filepath: str, delay_seconds: int = 15):
    """延迟删除文件（shred 级）"""
    def _del():
        time.sleep(delay_seconds)
        try:
            if os.path.exists(filepath):
                # shred: 覆写后删除
                with open(filepath, 'wb') as f:
                    f.write(os.urandom(min(os.path.getsize(filepath), 4096)))
                os.unlink(filepath)
                audit.log_file_deletion(os.path.basename(filepath), "service_end")
        except Exception:
            pass

    t = threading.Thread(target=_del, daemon=True)
    t.start()


def _generate_r1_questions(contract_summary):
    """LLM 根据合同内容生成第一轮 3 个问题。合同内容必须已脱敏。"""
    doc = contract_summary.get("doc_preview", "")[:4000]
    prompt = f"""你是一个商业合同分析系统的前端画像引擎。用户刚刚上传了一份合同，系统已经读取了合同全文。

请根据合同内容，生成 3 个针对性的画像问题。这些问题必须：
1. 与这份具体合同的内容直接相关，不能是泛泛而谈
2. 每个问题是封闭式选择题（4-6 个选项），选项之间互斥
3. Q1 必须是：询问用户在合同中的身份——是合同中的哪一方（甲方/乙方/转让方/受让方/收购方/被收购方/委托方/受托方等），根据合同内容合理推断可能的缔约方身份来生成选项，禁止在任何选项中列出合同出现的具体公司名称或自然人姓名
4. Q2-Q3 覆盖维度：交易核心利益焦点、最大风险感知
5. 语言锋锐、专业、简洁
6. 【安全红线】禁止在问题或选项中提及任何具体公司名称、自然人姓名、精确金额数字或地址

合同基本信息：
- 文件名：{contract_summary.get('filename', '')}
- 条款数：{contract_summary.get('clause_count', 0)}

合同内容片段（已脱敏处理，请基于此回答问题）：
{doc}

请输出严格 JSON：
{{"questions": [{{"id": "Q1", "title": "<=12字标题", "question": "完整问题", "options": ["选项A", "选项B", "选项C", "选项D"]}}, ...]}}"""

    try:
        resp = llm.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[{"role": "system", "content": "你是乔曦，商业合同分析助理。输出严格 JSON，不要 Markdown 包裹。禁止在选项中列出具体公司名或人名。"},
                      {"role": "user", "content": prompt}],
            temperature=0.7, max_tokens=2048, response_format={"type": "json_object"},
        )
        return json.loads(resp.choices[0].message.content).get("questions", [])
    except Exception:
        return [
            {"id": "Q1", "title": "合同身份", "question": "您在本合同中的身份是？",
             "options": ["甲方/收购方", "乙方/出售方", "丙方/担保方", "其他缔约方"]},
            {"id": "Q2", "title": "交易核心", "question": "本次交易中您最看重什么？",
             "options": ["交易对价的合理性", "标的权属的清晰度", "交割时间与节奏", "后续经营控制权"]},
            {"id": "Q3", "title": "最大风险", "question": "您对这份合同最大的担忧是什么？",
             "options": ["标的真实性存疑", "资金安全/付款后无法确权", "税务/合规风险", "交割后失控/退出路径不清晰"]},
        ]


def _generate_r2_questions(contract_summary, r1_answers, r1_questions):
    r1_text = "\n".join([
        "%s: %s -> %s" % (q.get("title", ""), q.get("question", ""), r1_answers.get(q.get("id", "")))
        for q in r1_questions
    ])
    doc = contract_summary.get("doc_preview", "")[:3000]
    prompt = f"""你是一个商业合同分析系统的前端画像引擎。以下是用户的第一轮画像回答：

{r1_text}

合同内容片段（已脱敏）：
{doc}

请基于用户的回答和合同内容，生成 3 个追问。追问维度：
1. 风险偏好与最大损失容忍度
2. 谈判地位与替代方案（BATNA）
3. 可妥协条件与绝对不可退让的底线

要求：
- 与用户的具体回答和合同内容紧密相关，不能泛泛
- 每个问题 4-5 个互斥选项
- 禁止在选项中出现具体公司名、人名、精确金额
- 语言锋锐专业

输出严格 JSON：
{{"questions": [{{"id": "Q4", "title": "<=12字", "question": "完整问题", "options": [...]}}, ...]}}"""

    try:
        resp = llm.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[{"role": "system", "content": "你是乔曦。输出严格 JSON。禁止列出具体公司名/人名。"},
                      {"role": "user", "content": prompt}],
            temperature=0.7, max_tokens=2048, response_format={"type": "json_object"},
        )
        return json.loads(resp.choices[0].message.content).get("questions", [])
    except Exception:
        return [
            {"id": "Q4", "title": "风险偏好", "question": "您的风险承受能力？", "options": ["保守", "适中", "激进"]},
            {"id": "Q5", "title": "谈判地位与BATNA", "question": "您的处境是？",
             "options": ["买家强势有替代", "买家弱势稀缺", "双方均势", "我是卖方"]},
            {"id": "Q6", "title": "绝对底线", "question": "您最不能退让的是什么？",
             "options": ["禁止先付款后确权", "财税控制权不可转让", "不接受单方核弹级违约", "对手必须提供资质保证"]},
        ]


def _parse_profile_from_answers(answers_text):
    try:
        resp = llm.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[{"role": "system", "content": """解析用户画像回答为 JSON。
anxiety_focus: authenticity|funds|tax|control|exit|compliance
risk_appetite: conservative|moderate|aggressive
max_loss_pct: 0.05|0.15|0.30
position: buyer_strong|buyer_weak|equal|seller|cooperator
batna_strength: strong|weak|none
hard_lines: no_prepayment_without_guarantee|fiscal_control_untransferable|no_unilateral_nuke
compromise_dims: price|schedule|governance|scope
输出严格 JSON。"""},
                      {"role": "user", "content": answers_text}],
            temperature=0.2, max_tokens=1024, response_format={"type": "json_object"},
        )
        parsed = json.loads(resp.choices[0].message.content)
        sl = parsed.get("strategic_layer", {})
        tl = parsed.get("tactical_layer", {})
        return {
            "client_id": "QX-%s" % datetime.now().strftime("%Y%m%d%H%M%S"),
            "upload_timestamp": datetime.now().isoformat(),
            "basic_info": {"client_name": "", "industry": parsed.get("industry", "other"),
                           "counterparty_name": "", "transaction_type": parsed.get("transaction_type", "other"),
                           "contract_value_cny": 0, "contract_duration_months": 0},
            "strategic_layer": {
                "interest_weights": sl.get("interest_weights", {"control": 0.4, "cashflow": 0.3, "tax": 0.15, "time": 0.15}),
                "trauma_tags": sl.get("trauma_tags", []),
                "anxiety_focus": sl.get("anxiety_focus", "authenticity"),
            },
            "tactical_layer": {
                "risk_appetite": tl.get("risk_appetite", "moderate"),
                "max_loss_pct": tl.get("max_loss_pct", 0.15),
                "position": tl.get("position", "equal"),
                "batna_strength": tl.get("batna_strength", "weak"),
                "compromise_dims": tl.get("compromise_dims", []),
                "hard_lines": tl.get("hard_lines", []),
            },
            "system_flags": {"open_input_used": False, "handoff_triggered": False, "confidence_score": 1.0},
        }
    except Exception:
        from src.state0_profile import build_client_profile
        return build_client_profile(
            {"client_name": "", "industry": "other", "counterparty_name": "", "transaction_type": "other",
             "contract_value_cny": 0, "contract_duration_months": 0},
            {"interest_weights": {"control": 0.4, "cashflow": 0.3, "tax": 0.15, "time": 0.15},
             "trauma_tags": [], "anxiety_focus": "authenticity"},
            {"risk_appetite": "moderate", "max_loss_pct": 0.15, "position": "equal",
             "batna_strength": "weak", "compromise_dims": [], "hard_lines": []},
        )


# ═══════════════════════════════════════════════════════════════
# 错误边界
# ═══════════════════════════════════════════════════════════════
if st.session_state.error_msg:
    st.error("系统错误: %s" % st.session_state.error_msg)
    if st.button("🔄 重新开始"):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()
    st.stop()
# ╔══════════════════════════════════════════════════════════════╗
# ║  STEP -1: 付款验证                                           ║
# ╚══════════════════════════════════════════════════════════════╝
_phase = st.session_state.phase
if _phase == "payment":

    st.markdown("### 💳 服务付款")
    st.markdown("使用 **霖信莯 · Qiaoxi 商业合同审查系统** 需支付基础服务费。付款后获得本次审查服务的使用授权码，输入后即可进入系统。")

    st.divider()

    col_info, col_qr = st.columns([3, 2])

    with col_info:
        st.markdown("""
**基础服务包含：**
- ✅ 合同全文结构化解析
- ✅ 法律风险全面标定（含法规 RAG 检索）
- ✅ 商业结构与系统动力学分析
- ✅ 六位独立评审员并行审计
- ✅ Qiaoxi Contract-Analyzer 最终决策报告
- ✅ 完整《商业决策报告》下载（Markdown + Word）

**收费标准：¥20 / 次**（单份合同，单次审查）

付款方式：扫描右侧收款码，备注"合同审查"，付款后截图发送至客服微信获取授权码。
""")
        st.info("📞 联系人：余磊 13987671259　｜　📧 邮箱：425448719@qq.com")

    with col_qr:
        st.markdown("**扫码付款（微信 / 支付宝）**")
        # ── 收款码占位区 ──
        # 将收款码图片放置于 assets/qr_payment.png 后取消下方注释：
        # st.image("assets/qr_payment.png", width=200)
        st.markdown(
            """<div style="width:200px;height:200px;border:2px dashed #aaa;border-radius:8px;
            display:flex;align-items:center;justify-content:center;color:#aaa;font-size:13px;
            text-align:center;padding:12px;">
            收款码<br/>待上传<br/><br/>assets/qr_payment.png
            </div>""",
            unsafe_allow_html=True,
        )

    st.divider()
    st.markdown("#### 输入授权码")
    col_code, col_btn = st.columns([3, 1])
    pay_code_input = col_code.text_input(
        "授权码", key="pay_code_input",
        placeholder="QIAOXI-XXXX-XXXX 或 CXL-XXXX-XXXX",
        label_visibility="collapsed",
    )
    if col_btn.button("验证并进入", type="primary", use_container_width=True):
        code_input = pay_code_input.strip().upper()

        # ── 先验证网站体验码 ──
        if code_input.startswith("CXL-"):
            try:
                resp = _requests.post(
                    "http://localhost:3000/api/trial/verify",
                    json={"code": code_input},
                    timeout=5
                )
                data = resp.json()
                if data.get("valid") and data.get("remaining", {}).get("qiaoxi", 0) > 0:
                    st.session_state.base_paid = True
                    st.session_state.trial_code = code_input
                    st.session_state.phase = "consent"
                    st.rerun()
                elif data.get("valid"):
                    st.error("该体验码的乔曦额度已用完，请使用付费授权码。")
                else:
                    st.error(data.get("error", "体验码无效。"))
            except Exception:
                st.error("体验码验证失败，请检查网络连接或使用付费授权码。")
        # ── 原有授权码验证 ──
        else:
            VALID_BASE_CODES = {"QIAOXI-DEMO-2024", "QIAOXI-BETA-0001", "QIAOXI-PAY-0001"}
            if code_input in VALID_BASE_CODES:
                st.session_state.base_paid = True
                st.session_state.trial_code = None
                st.session_state.phase = "consent"
                st.rerun()
            else:
                st.error("授权码无效。请确认付款已完成并联系客服获取正确的授权码。")

    st.caption("如需发票或企业批量采购报价，请联系霖信莯咨询。")


# ╔══════════════════════════════════════════════════════════════╗
# ║  STEP 0: 授权协议                                           ║
# ╚══════════════════════════════════════════════════════════════╝
elif _phase == "consent":

    if not st.session_state.base_paid:
        st.session_state.phase = "payment"
        st.rerun()

    st.markdown("### 📜 用户授权与服务协议")
    st.caption("请仔细阅读以下协议。您必须同意所有条款才能使用本系统。")

    st.markdown('<div class="consent-box">%s</div>' %
                USER_CONSENT_TEMPLATE.replace('\n', '<br/>'), unsafe_allow_html=True)

    st.divider()

    c1 = st.checkbox("我确认已阅读并理解上述数据安全措施", key="consent_security")
    c2 = st.checkbox("我确认已知悉本系统为商业决策辅助工具，不构成正式法律意见", key="consent_disclaimer")
    c3 = st.checkbox("我授权霖信莯在本地服务器上处理我的合同文件，霖信莯 · Qiaoxi Contract-Analyzer 仅接收脱敏后的信息", key="consent_authorize")

    if c1 and c2 and c3:
        st.success("✅ 您已同意所有条款。请点击下方按钮进入系统。")
        if st.button("进入系统 →", type="primary", use_container_width=True):
            st.session_state.consent_granted = True
            audit.log_user_consent(True)
            st.session_state.phase = "upload"
            st.rerun()
    else:
        st.info("请勾选全部三个确认项以继续。")


# ╔══════════════════════════════════════════════════════════════╗
# ║  STEP 1: 上传合同                                           ║
# ╚══════════════════════════════════════════════════════════════╝
elif _phase == "upload":

    if not st.session_state.consent_granted:
        st.warning("请先完成授权协议。")
        st.stop()

    st.markdown("### 📄 第一步：上传合同文件")
    st.caption("支持 PDF / DOCX。系统将先对合同执行脱敏，再让霖信莯 · Qiaoxi Contract-Analyzer 系统的分析引擎读取脱敏后的内容。")

    uploaded_file = st.file_uploader("拖拽文件到此处，或点击浏览", type=["pdf", "docx"], key="contract_upload")
    do_parse = st.button("📤 上传并开始脱敏分析", type="primary", use_container_width=True)

    if do_parse and uploaded_file is None:
        st.warning("请先上传一份合同文件。")

    if do_parse and uploaded_file is not None:
        try:
            ext = os.path.splitext(uploaded_file.name)[1].lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                tmp.write(uploaded_file.getvalue())
                st.session_state.uploaded_file_path = tmp.name
                st.session_state.uploaded_file = uploaded_file
                st.session_state.session_files.append(tmp.name)

            with st.spinner("Step 1/2: 正在解析合同原始文本..."):
                parser = QiaoxiContractParser()
                result = parser.parse_contract(st.session_state.uploaded_file_path, rag_results=None)
                if "error" in result:
                    st.error("合同解析失败: " + result["error"]); st.stop()
                st.session_state.contract_raw = result.get("document_text", "")

            with st.spinner("Step 2/2: 正在脱敏处理（分析引擎仅能看见脱敏后的版本）..."):
                sanitized = deep_sanitize_contract_text(st.session_state.contract_raw or "")
                st.session_state.contract_sanitized = sanitized
                # 用脱敏文本重新提取条款树
                st.session_state.clause_tree = parser._extract_clause_tree(sanitized)

                st.session_state.contract_summary = {
                    "filename": uploaded_file.name,
                    "clause_count": st.session_state.clause_tree.get("total_clauses", 0),
                    "doc_preview": sanitized[:4000],
                }

            audit.log_state_transition("0_consent_upload", "0_pre_read_sanitized")
            st.session_state.phase = "profile_r1"
            st.rerun()
        except Exception as e:
            st.error("上传/脱敏阶段出错: %s" % str(e))
            st.stop()


# ╔══════════════════════════════════════════════════════════════╗
# ║  STEP 2: 第一轮画像                                         ║
# ╚══════════════════════════════════════════════════════════════╝
elif _phase == "profile_r1":

    st.markdown("### 📋 第二步：客户画像 · 第一轮")
    st.caption("以下问题是霖信莯 · Qiaoxi Contract-Analyzer 根据脱敏后的合同内容针对性生成的。分析引擎无法看到具体公司名或人名。")

    try:
        if st.session_state.r1_questions is None:
            with st.spinner("霖信莯 · Qiaoxi Contract-Analyzer 正在根据脱敏合同内容生成针对性问题..."):
                st.session_state.r1_questions = _generate_r1_questions(st.session_state.contract_summary)

        r1q = st.session_state.r1_questions
        r1a = {}
        for qi, q in enumerate(r1q):
            qid = q.get("id", "?")
            opts = q.get("options", [])
            is_q1 = (qid == "Q1" or qi == 0)
            st.markdown('<div class="q-card"><strong>%s：%s</strong><br/>%s</div>' % (
                qid, q.get("title", ""), q.get("question", "")), unsafe_allow_html=True)

            if is_q1:
                # Q1 是合同身份确认，只能单选（用 radio 按钮）
                ans = st.selectbox("请选择您在本合同中的身份（单选）", opts, key="r1sel_%s" % qid, index=None,
                                   placeholder="请选择...")
                r1a[qid] = [ans] if ans else []
            else:
                selected = []
                max_sel = min(3, len(opts))
                st.caption("可多选，至多选%d项" % max_sel)
                for i, opt in enumerate(opts):
                    # 当前已选 >= 上限时，其余 checkbox disabled
                    already_reached = len(selected) >= max_sel
                    if st.checkbox(opt, key="r1cb_%s_%d" % (qid, i), disabled=already_reached and opt not in selected):
                        selected.append(opt)
                r1a[qid] = selected

        st.divider()
        # 进度提示条（做完题可见）
        st.info("📊 点击下方按钮，系统将进入追问环节")
        st.progress(0.33, text="进度：第一轮画像完成 → 等待进入追问")

        do_r1_bottom = st.button("✅ 提交 → 进入追问", type="primary", use_container_width=True, key="do_r1_bottom")
        if do_r1_bottom:
            unanswered = [q.get("id") for q in r1q if not r1a.get(q.get("id"))]
            if unanswered:
                st.error("请回答所有问题。未回答: %s" % ", ".join(unanswered))
                st.stop()
            st.session_state.r1_answers = r1a
            st.session_state.phase = "profile_r2"
            st.rerun()
    except Exception as e:
        st.error("第一轮画像出错: %s" % str(e))
        st.stop()


# ╔══════════════════════════════════════════════════════════════╗
# ║  STEP 3: 第二轮追问                                         ║
# ╚══════════════════════════════════════════════════════════════╝
elif _phase == "profile_r2":

    st.markdown("### 📋 第三步：客户画像 · 追问")
    st.caption("以下问题是系统根据第一轮回答和脱敏合同内容动态生成的。")

    try:
        if st.session_state.r2_questions is None:
            with st.spinner("系统正在根据您的回答生成追问..."):
                st.session_state.r2_questions = _generate_r2_questions(
                    st.session_state.contract_summary,
                    st.session_state.r1_answers,
                    st.session_state.r1_questions,
                )

        r2q = st.session_state.r2_questions
        r2a = {}
        for qi, q in enumerate(r2q):
            qid = q.get("id", "?")
            opts = q.get("options", [])
            is_q4 = (qid == "Q4" or qi == 0)
            st.markdown('<div class="q-card"><strong>%s：%s</strong><br/>%s</div>' % (
                qid, q.get("title", ""), q.get("question", "")), unsafe_allow_html=True)

            if is_q4:
                # Q4 是风险偏好，只能单选
                ans = st.selectbox("请选择（单选）", opts, key="r2sel_%s" % qid, index=None,
                                   placeholder="请选择...")
                r2a[qid] = [ans] if ans else []
            else:
                selected = []
                max_sel = min(3, len(opts))
                st.caption("可多选，至多选%d项" % max_sel)
                for i, opt in enumerate(opts):
                    already_reached = len(selected) >= max_sel
                    if st.checkbox(opt, key="r2cb_%s_%d" % (qid, i), disabled=already_reached and opt not in selected):
                        selected.append(opt)
                r2a[qid] = selected

        st.divider()
        # 进度提示条（第二轮）
        st.success("✅ 所有问题已回答完毕，系统准备就绪")
        st.info("📊 点击下方按钮，系统将开始分析。请耐心等待，不要关闭页面，处理可能需要一点时间。")
        st.progress(0.66, text="进度：两轮画像完成 → 等待进入审查")

        do_r2_bottom = st.button("✅ 确认画像 → 开始审查", type="primary", use_container_width=True, key="do_r2_bottom")
        if do_r2_bottom:
            unanswered = [q.get("id") for q in r2q if not r2a.get(q.get("id"))]
            if unanswered:
                st.error("请回答所有问题。未回答: %s" % ", ".join(unanswered))
                st.stop()

            st.markdown("""
<div style="background:#e8f4fd;border:2px solid #1890ff;border-radius:10px;padding:20px 24px;margin:16px 0;text-align:center;">
  <div style="font-size:22px;margin-bottom:8px;">⏳</div>
  <div style="font-size:16px;font-weight:700;color:#0f3460;margin-bottom:6px;">
    Qiaoxi Contract-Analyzer 系统正在处理您的信息需求
  </div>
  <div style="font-size:14px;color:#555;line-height:1.8;">
    请耐心等待，系统正在调用法规数据库并分析合同内容，预计需要 30–60 秒。<br/>
    <strong style="color:#cf1322;">请不要刷新页面，谢谢您的耐心。</strong><br/>
    您也可以回到本页面上方查看业务处理的进度条。
  </div>
</div>
""", unsafe_allow_html=True)

            # 点击开始审查后，显示等待提示
            with st.spinner("请耐心等待，不要关闭页面，系统正在处理信息和分析问题，这可能需要花点时间..."):
                import time
                time.sleep(0.5)  # 让 spinner 显示出来

                st.session_state.r2_answers = r2a

                all_answers_text = ""
                for q in st.session_state.r1_questions:
                    all_answers_text += "Q: %s\nA: %s\n" % (
                        q.get("question", ""), st.session_state.r1_answers.get(q.get("id", "")))
                for q in st.session_state.r2_questions:
                    all_answers_text += "Q: %s\nA: %s\n" % (
                        q.get("question", ""), r2a.get(q.get("id", "")))

                st.session_state.profile = _parse_profile_from_answers(all_answers_text)

            st.session_state.phase = "processing"
            st.rerun()
    except Exception as e:
        st.error("第二轮追问出错: %s" % str(e))
        st.stop()


# ╔══════════════════════════════════════════════════════════════╗
# ║  STEP 4: 审查处理                                          ║
# ╚══════════════════════════════════════════════════════════════╝
elif _phase == "processing":

    progress_bar = st.progress(0, text="准备中...")
    status_text = st.empty()

    try:
        progress_bar.progress(15, text="步骤 1/6: 合同条款树（已使用脱敏版本）")
        if st.session_state.clause_tree is None:
            parser = QiaoxiContractParser()
            sanitized = st.session_state.contract_sanitized or ""
            st.session_state.clause_tree = parser._extract_clause_tree(sanitized)

        progress_bar.progress(25, text="步骤 1/6: 条款树提取完成")

        progress_bar.progress(33, text="步骤 2/6: 乔曦法律初审中（审查脱敏条款）...")
        status_text.info("乔曦正在进行法律风险标定（锋锐模态）...")
        audit.log_state_transition("1_parse", "2_legal_review")

        # RAG 法规检索
        from src.rag_retriever import search_relevant_laws
        sanitized_text = st.session_state.contract_sanitized or st.session_state.contract_raw or ""
        rag_results = search_relevant_laws(sanitized_text, top_k=5)
        audit.log_rag_call(success=len(rag_results) > 0, query_count=len(rag_results))

        reviewer = QiaoxiLegalReviewer()
        # 乔曦审查脱敏后的条款，代入客户画像 + RAG 法规检索结果
        legal_review = reviewer.review(
            st.session_state.clause_tree,
            profile=st.session_state.profile,
            rag_results=rag_results,
        )
        st.session_state.legal_review = legal_review
        progress_bar.progress(50, text="步骤 2/6: 法律初审完成（已代入客户立场 + 法规数据库检索 + 法律分析模型）")

        # ─── State 3: 商业模式解构（CLD）───
        progress_bar.progress(55, text="步骤 3/6: 商业模式系统动力学建模...")
        status_text.info("正在提取资金流向、权力分配、时间约束...")
        try:
            cld_builder = CLDBuilder()
            cld_report = cld_builder.build(
                st.session_state.clause_tree,
                legal_review=legal_review,
            )
            st.session_state.cld_report = cld_report
            audit.log_state_transition("2_legal_review", "3_cld")
            progress_bar.progress(62, text="步骤 3/6: 商业模式解构完成（已提取因果回路 + 关键变量）")
        except Exception as e:
            logger.error(f"[State 3 CLD] 失败: {e}")
            st.session_state.cld_report = {"loops": [], "key_variables": {}, "summary": "解构失败", "error": str(e)}
            progress_bar.progress(62, text="步骤 3/6: 商业模式解构失败（将使用兜底数据）")

        # ─── State 4: 六位评审员串行审计 ───
        progress_bar.progress(64, text="步骤 4/6: 私董会六位评审员审计中...")
        status_text.info("六位评审员串行审查中，每人独立评估，请耐心等待...")
        try:
            council_runner = CouncilRunner()

            # 前端动态进度显示
            auditor_progress = st.empty()
            def _council_progress(idx: int, name: str):
                auditor_progress.info(f"六位评审员 {idx}/6 {name} 正在审查中...")

            council_result = council_runner.run_council(
                clause_tree=st.session_state.clause_tree,
                legal_review=legal_review,
                cld_report=st.session_state.cld_report,
                progress_callback=_council_progress,
            )
            st.session_state.six_audits = council_result["audits"]
            st.session_state.veto_any = council_result["veto_any"]
            st.session_state.veto_auditors = council_result["veto_auditors"]
            audit.log_state_transition("3_cld", "4_council")
            auditor_progress.success(f"六位评审员 6/6 全部完成（否决: {len(council_result['veto_auditors'])}/6）")
            progress_bar.progress(72, text="步骤 4/6: 私董会审计完成")
            if council_result["veto_any"]:
                status_text.warning(f"⚠️ 否决触发: {'、'.join(council_result['veto_auditors'])}")
        except Exception as e:
            logger.error(f"[State 4 Council] 失败: {e}")
            st.session_state.six_audits = []
            st.session_state.veto_any = False
            st.session_state.veto_auditors = []
            progress_bar.progress(72, text="步骤 4/6: 私董会审计失败（将使用兜底数据）")

        # ─── State 5: 推演引擎 + 辩论合成 ───
        progress_bar.progress(74, text="步骤 5/6: 时间轴推演 + 私董会辩论合成...")
        status_text.info("推演引擎正在模拟 3/6/12/36 月时间切片...")
        try:
            # 确定性推演
            sim_engine = SimulationEngine()
            simulation_result = sim_engine.run(
                cld_report=st.session_state.cld_report,
                audits=st.session_state.six_audits,
            )
            st.session_state.simulation = simulation_result

            # LLM 辩论合成
            debate_synth = DebateSynthesizer()
            debate_result = debate_synth.synthesize(
                audits=st.session_state.six_audits,
                simulation=simulation_result,
                cld_report=st.session_state.cld_report,
            )
            st.session_state.debate = debate_result
            audit.log_state_transition("4_council", "5_simulation")
            progress_bar.progress(80, text=f"步骤 5/6: 推演完成（轨迹: {simulation_result.get('trajectory', '未知')}）")
        except Exception as e:
            logger.error(f"[State 5 Simulation] 失败: {e}")
            st.session_state.simulation = {"snapshots": [], "trajectory": "未知", "error": str(e)}
            st.session_state.debate = {"debate_minutes": "推演失败"}
            progress_bar.progress(80, text="步骤 5/6: 推演失败（将使用兜底数据）")

        # ─── State 6: 李超逸决策 ───
        progress_bar.progress(83, text="步骤 6/6: 李超逸最终决策中...")
        status_text.info("李超逸正在综合所有信息做出最终决策...")
        try:
            decision_engine = DecisionEngine()
            decision_result = decision_engine.decide(
                audits=st.session_state.six_audits,
                simulation=st.session_state.simulation,
                cld_report=st.session_state.cld_report,
                profile=st.session_state.profile,
                legal_review=legal_review,
            )
            st.session_state.decision = decision_result
            audit.log_state_transition("5_simulation", "6_decision")

            decision_label = decision_result.get("decision_label", "?")
            if decision_result.get("veto_triggered"):
                status_text.error(f"🛑 李超逸最终决策: {decision_label} (否决触发)")
            else:
                status_text.success(f"李超逸最终决策: {decision_label}")
            progress_bar.progress(90, text=f"步骤 6/6: 决策完成 → {decision_label}")
        except Exception as e:
            logger.error(f"[State 6 Decision] 失败: {e}")
            st.session_state.decision = {"decision_label": "拖", "veto_triggered": False, "handoff_required": True, "detailed_order": f"决策引擎故障: {e}"}
            progress_bar.progress(90, text="步骤 6/6: 决策引擎故障（已触发人工介入建议）")

        progress_bar.progress(95, text="报告生成中...")

        # ─── 报告生成（客户视角，不暴露内部流程/人名/代号）───
        risks = legal_review.get("risks", [])
        high_risks = [r for r in risks if r.get("risk_level") == "high"]
        mid_risks = [r for r in risks if r.get("risk_level") == "medium"]
        low_risks = [r for r in risks if r.get("risk_level") == "low"]
        profile = st.session_state.profile or {}
        sl = profile.get("strategic_layer", {})
        tl = profile.get("tactical_layer", {})
        fn = st.session_state.uploaded_file.name if st.session_state.uploaded_file else "N/A"

        # State 3-6 数据
        cld_report = st.session_state.cld_report or {}
        six_audits = st.session_state.six_audits or []
        simulation = st.session_state.simulation or {}
        debate = st.session_state.debate or {}
        decision = st.session_state.decision or {}

        # ─── 辅助函数 ───
        INTERNAL_NAMES = ["李文鸿", "吴慧琼", "李军", "段海涛", "王志坚", "李艾熹", "李超逸", "乔曦",
                          "value_investor", "cfo_risk", "industry_architect", "deal_engineer",
                          "operations", "risk_philosopher"]
        ANGLE_MAP = {
            "李文鸿": "价值投资人视角", "吴慧琼": "首席风控视角", "李军": "行业架构师视角",
            "段海涛": "交易结构视角", "王志坚": "运营落地视角", "李艾熹": "风险哲学视角",
        }

        def _clean_internal_names(text: str) -> str:
            """清洗文本中的内部角色名和 English 角色 key"""
            for kw in INTERNAL_NAMES:
                text = text.replace(kw, "")
            return text

        def _clause_label(clause_id: str) -> str:
            """将 CLS-0004 转为可读编号：第4条"""
            import re
            m = re.search(r'(\d+)', clause_id or "")
            if m:
                return f"第{m.group(1).lstrip('0') or '0'}条"
            return clause_id or "某条款"

        POSITION_MAP = {"buyer_strong": "买方强势，有多种替代方案",
                        "buyer_weak": "买方弱势，标的稀缺，可选空间小",
                        "equal": "双方均势，有谈判空间",
                        "seller": "作为卖方",
                        "cooperator": "合作方身份"}
        ANXIETY_MAP = {"authenticity": "标的真实性——合同涉及的核心资产或权利是否真实存在、权属清晰",
                       "funds": "资金安全——付款后能否锁定相应权利",
                       "tax": "税务合规——交易涉及的税务敞口与合规风险",
                       "control": "控制权——能否在交割后真正掌控标的资产或公司",
                       "exit": "退出路径——如果交易失败，已投入资金能否收回",
                       "compliance": "合规风险——交易是否触碰监管红线"}
        RISK_MAP = {"conservative": "保守——宁愿少赚，不能亏损",
                    "moderate": "适中——接受可控风险换取合理回报",
                    "aggressive": "积极——愿意承担较高风险博取高回报"}
        BATNA_MAP = {"strong": "坚强——当前已有替代交易方案",
                     "weak": "较弱——替代方案不理想，但仍有其他选择",
                     "none": "无替代方案——本条交易是唯一选项"}

        report = f"""# 合同分析报告

> **报告日期**: {datetime.now().strftime('%Y年%m月%d日')}
> **合同文件**: {fn}
> **出具方**: 霖信莯信息咨询
> **隐私声明**: 本报告基于脱敏处理后的合同内容生成。霖信莯 Qiaoxi 分析引擎在分析过程中仅接触脱敏后的信息，未获取原始合同中的具体公司名称、自然人名、证件号或精确金额。

---

## 一、合同概览

"""
        # ─── 第一章：合同概览 ───
        # 用原始脱敏文本重新精确计数"第X条"，比 clause_tree 的正则第一轮匹配更可靠
        sanitized_text = st.session_state.contract_sanitized or ""
        clause_count = len(_re.findall(r'第[一二三四五六七八九十百千\d]+条', sanitized_text))
        if clause_count == 0:
            # 回退到 clause_tree 的记录
            clause_count = st.session_state.clause_tree.get('total_clauses', 0) if st.session_state.clause_tree else 0
        clauses_list = st.session_state.clause_tree.get('clauses', []) if st.session_state.clause_tree else []

        # 提取合同标题：优先用文件名，比"第一条"更合理
        title = st.session_state.contract_summary.get("filename", "未提取") if st.session_state.contract_summary else "未提取"
        # 去掉文件扩展名
        if title != "未提取" and "." in title:
            title = title.rsplit(".", 1)[0]

        cld_summary = cld_report.get("summary", "")
        report += f"""- **合同名称**: {title}
- **分析范围**: 本报告覆盖合同条款文本分析、法律风险标定、商业结构解构和风险推演
"""
        # 交易结构一句话（从 CLD 摘要中提取）
        if cld_summary and len(cld_summary) > 20:
            report += f"""
**交易结构摘要**:
{cld_summary[:400]}
"""
        else:
            report += "\n> 交易结构分析未生成，可能因合同文本过短或解析失败。\n"

        # ─── 第二章：您的核心关切 ───
        position_cn = POSITION_MAP.get(tl.get('position', ''), tl.get('position', '未设置'))
        anxiety_cn = ANXIETY_MAP.get(sl.get('anxiety_focus', ''), sl.get('anxiety_focus', '未设置'))
        risk_cn = RISK_MAP.get(tl.get('risk_appetite', ''), tl.get('risk_appetite', '未设置'))
        batna_cn = BATNA_MAP.get(tl.get('batna_strength', ''), tl.get('batna_strength', '未设置'))
        max_loss = tl.get('max_loss_pct', 0) * 100
        hard_lines = tl.get('hard_lines', [])
        compromise = tl.get('compromise_dims', [])

        # 底线中文映射
        HARD_LINE_MAP = {
            "no_prepayment_without_guarantee": "禁止先付款后确权（必须先锁定权利再付款）",
            "fiscal_control_untransferable": "财税控制权不可转让",
            "no_unilateral_nuke": "不接受单方核弹级违约条款",
        }
        COMPROMISE_MAP = {
            "price": "价格（可适当调整对价）",
            "schedule": "时间（可调整交割节奏）",
            "governance": "治理（可协商公司治理结构）",
            "scope": "范围（可调整交易范围）",
        }
        hard_cn = [HARD_LINE_MAP.get(h, h) for h in hard_lines] if hard_lines else []
        comp_cn = [COMPROMISE_MAP.get(c, c) for c in compromise] if compromise else []

        report += f"""
## 二、您的核心关切

根据您在画像环节的选择，您在本次交易中的核心立场和风险偏好如下：

| 维度 | 您的选择 |
|------|---------|
| 交易地位 | {position_cn} |
| 最担心的问题 | {anxiety_cn} |
| 风险偏好 | {risk_cn} |
| 最多可接受的损失 | 不超过净资产的 {max_loss:.0f}% |
| 替代方案 | {batna_cn} |
"""
        if hard_cn:
            report += f"| 绝对不可退让的底线 | {'、'.join(hard_cn)} |\n"
        else:
            report += "| 绝对不可退让的底线 | 未明确设置（建议明确）|\n"
        if comp_cn:
            report += f"| 可以妥协的维度 | {'、'.join(comp_cn)} |\n"

        # 底线冲突检测
        violations = legal_review.get('bottom_line_violations', [])
        report += f"""
### 底线冲突检测
"""
        if violations:
            for v in violations:
                report += f"> ⛔ **底线冲突**: {v}\n\n"
        else:
            report += "> ✅ 本轮审查未检测到与您设定的绝对底线存在直接冲突。\n"

        # ─── 第三章：交易结构分析（通俗解读版）───
        cld_loops = cld_report.get("loops", [])
        cld_vars = cld_report.get("key_variables", {})

        report += f"""
## 三、交易结构分析

> **说明**：以下内容来自对合同交易结构的系统化拆解。我们使用了专业的商业模型分析工具，但已将分析结论翻译为您能读懂的语言。简单来说，这一章回答三个问题：**钱怎么走、权怎么移、风险怎么变**。

"""
        # CLD 摘要作为交易结构概述
        if cld_summary and len(cld_summary) > 20:
            report += f"**一句话看懂这笔交易**：{cld_summary}\n\n"

        # 用"如果…那么…"替代因果回路图
        if cld_loops:
            report += '### 这笔交易里的"连锁反应"\n\n'
            report += '下面列出的各种因果链，就是合同里不同条款之间的联动关系。您可以把它理解为一种"蝴蝶效应"——改动合同里的一点，可能在其他地方引发连锁反应。\n\n'
            for loop in cld_loops:
                loop_desc = loop.get("description", "")
                loop_id = loop.get("id", "")
                is_r = "reinforc" in loop.get("type", "")
                if is_r:
                    # 增强回路 → 滚雪球效应
                    report += f"**{loop_id} · 放大效应（正向循环）**\n"
                    report += f"> 如果 {loop_desc[:300]}\n"
                    report += '> ⚠️ 简单说：这会像滚雪球一样越滚越大，一旦启动就很难停下来。如果您希望这个趋势发生，那是好事；如果不希望，就需要在源头上设置刹车。\n\n'
                else:
                    # 调节回路 → 刹车机制
                    report += f"**{loop_id} · 刹车机制（制约因素）**\n"
                    report += f"> 如果 {loop_desc[:300]}\n"
                    report += '> ℹ️ 简单说：这是天然的"刹车"，当事情发展到一定程度会自动踩刹车。但如果刹车太强，也可能让交易推进困难。\n\n'

        # 资金流向
        cashflows = cld_vars.get("cashflow", [])
        if cashflows:
            report += "### 资金流向：钱怎么走？\n\n"
            report += "下表记录的是交易中每一笔钱的走向和触发条件。关注两点：**您什么时候付钱**，以及**付钱的时候您拿到了什么**。\n\n"
            for c in cashflows[:6]:
                report += f"- {c}\n"
            report += "\n"

        # 权利分配
        powers = cld_vars.get("power", [])
        if powers:
            report += "### 权利与控制权：谁说了算？\n\n"
            report += "这部分展示的是交易完成后，谁掌握公章、谁当法人、谁有董事席位，以及这些权利什么时候从对方手里转移到您手里。\n\n"
            for p in powers[:6]:
                report += f"- {p}\n"
            report += "\n"

        # 时间线
        times = cld_vars.get("time", [])
        if times:
            report += "### 关键时间节点：什么时候该做什么？\n\n"
            report += "以下是交易的关键时间表。特别注意各个节点的先后顺序——顺序错了，风险就来了。\n\n"
            for t in times[:6]:
                report += f"- {t}\n"
            report += "\n"

        # 推演结果
        sim_traj = simulation.get("trajectory", "")
        sim_snapshots = simulation.get("snapshots", [])
        if sim_snapshots:
            report += "### 风险趋势推演：未来会怎样？\n\n"
            report += "以下为基于交易结构的推算——如果各方按照当前约定的节奏和条件推进，在不同时间点您的处境会是：\n\n"
            report += "| 时间 | 风险水平 | 在这个时点您可能面临什么 |\n"
            report += "|------|---------|--------|\n"
            for s in sim_snapshots[:4]:
                m = s.get("month", "?")
                level = s.get("risk_level", "?")
                level_cn = {"high": "🔴 高", "medium": "🟡 中", "low": "🟢 低"}.get(level, level)
                signals = s.get("risk_signals", [])
                signal_text = signals[0] if signals else "—"
                signal_text = signal_text.replace(f"M{m}: ", "")
                report += f"| 第 {m} 个月 | {level_cn} | {signal_text[:120]} |\n"

            if sim_traj:
                # 检查时间切片中的实际风险水平 — 如有多个高风险切片，则不应表述为"稳定"
                high_slices = sum(1 for s in sim_snapshots if s.get("risk_level") == "high")
                total_slices = len(sim_snapshots)
                traj_cn = {
                    "逐渐恶化": "⚠️ 如果保持现状，您的风险将随时间逐步恶化，建议尽快调整合同中的关键条款以扭转趋势",
                    "稳定向好": "✅ 随着关键节点完成，您的处境将逐步改善",
                    "风险稳定": {
                        0: "📊 风险水平预计保持稳定",
                        1: f"⚠️ 在 {total_slices} 个时间切片中有 {high_slices} 个处于高风险状态——虽然趋势没有恶化，但当前风险水平本身就偏高，建议重点关注",
                        2: f"⚠️ 在 {total_slices} 个时间切片中有 {high_slices} 个处于高风险状态——风险持续处于高位，虽未进一步恶化，但已处于危险区间",
                    }.get(high_slices, f"⚠️ 在 {total_slices} 个时间切片中有 {high_slices} 个处于高风险状态——风险持续高企，强烈建议采取措施降低风险暴露"),
                    "波动，拐点在前": "⚠️ 存在不确定因素，关键拐点在前方，需要密切监控"
                }.get(sim_traj, sim_traj)
                report += f"\n**总体判断**: {traj_cn}\n"
            report += "\n"

        # ─── 第四章：法律风险分析（乔曦初审，不露名）───
        report += f"""
## 四、法律风险分析

以下风险分析由霖信莯合同审查系统基于中国现行法律法规数据库完成。风险等级分为三级：高风险（需优先处理）、中风险（建议关注）、低风险（可酌情处理）。

"""
        # 高风险
        report += f"### 🔴 高风险（{len(high_risks)} 项）\n\n"
        if high_risks:
            for r in high_risks:
                cid = _clause_label(r.get('clause_id', '?'))
                desc = r.get('description', 'N/A')
                law = r.get('legal_basis', '【法规待核】')
                action = r.get('suggested_action', 'N/A')
                cat = r.get('risk_category', '未分类')
                report += f"""**{cid}** · {cat}

> 风险：{desc}
>
> 法律依据：{law}
>
> 建议：{action}

"""
        else:
            report += "> 未发现高风险项。\n\n"

        # 中风险
        report += f"### 🟡 中风险（{len(mid_risks)} 项）\n\n"
        if mid_risks:
            for r in mid_risks[:6]:
                cid = _clause_label(r.get('clause_id', '?'))
                desc = r.get('description', 'N/A')
                cat = r.get('risk_category', '未分类')
                report += f"- **{cid}**（{cat}）：{desc[:200]}\n"
            report += "\n"
        else:
            report += "> 未发现中风险项。\n\n"

        # 低风险
        report += f"### 🟢 低风险（{len(low_risks)} 项）\n\n"
        if low_risks:
            for r in low_risks[:4]:
                cid = _clause_label(r.get('clause_id', '?'))
                desc = r.get('description', 'N/A')
                report += f"- **{cid}**：{desc[:150]}\n"
            report += "\n"
        else:
            report += "> 未发现低风险项。\n\n"

        # 综合法律评估
        overall = legal_review.get('overall_client_assessment', '')
        if overall:
            report += f"""
### 综合法律评估

{overall}
"""

        # ─── 第五章：商业风险深度透视 ───
        report += f"""
## 五、商业风险深度透视

霖信莯内部采用了多视角交叉审查的方法，从资本安全、风控专业、行业经验、交易结构、运营落地和反面论证五个角度对合同进行了全面审查。以下是综合审查意见的汇总：

"""
        # 综合六位评审员输出为一篇连贯文本
        if six_audits:
            # 关键发现汇总
            high_concerns = []
            other_concerns = []
            for a in six_audits:
                summary = a.get("audit_summary", "")
                if a.get("veto_triggered"):
                    high_concerns.append(summary[:400])
                else:
                    other_concerns.append(summary[:300])

            if high_concerns:
                report += "### ⚠️ 核心风险警告\n\n"
                report += "经过多角度交叉审查，审查团队一致认为以下风险可能对您造成实质性损害：\n\n"
                for hc in high_concerns:
                    report += f"{_clean_internal_names(hc)}\n\n"

            # 共识标注
            consensus_items = debate.get("consensus_items", [])
            if consensus_items:
                report += "### 审查共识\n\n"
                report += "以下条款被多个审查视角独立识别为关键风险点：\n\n"
                for c in consensus_items:
                    label = _clause_label(c)
                    report += f"- **{label}**\n"
                report += "\n"

            # 少数观点
            minority = debate.get("minority_views", [])
            if minority:
                report += "### 其他值得关注的角度\n\n"
                for m in minority:
                    label = _clause_label(m.get("clause_id", ""))
                    who = m.get("identified_by", [])
                    # 人名 → 视角标签，不向客户暴露六位评审员个人姓名
                    ANGLE_MAP = {
                        "李文鸿": "价值投资人视角", "吴慧琼": "首席风控视角", "李军": "行业架构师视角",
                        "段海涛": "交易结构视角", "王志坚": "运营落地视角", "李艾熹": "风险哲学视角",
                    }
                    who_clean = [ANGLE_MAP.get(name, name) for name in who]
                    who_str = "、".join(who_clean)
                    report += f"- {label}：{who_str}等审查视角提出了额外关注\n"
                report += "\n"

        # ─── 第六章：决策建议 ───
        decision_label = decision.get("decision_label", "未生成")
        decision_reasons = decision.get("decision_reasons", [])
        detailed_order = decision.get("detailed_order", "")

        # 清理 internal 内容
        if "李超逸" in detailed_order:
            detailed_order = detailed_order.replace("李超逸", "霖信莯")
        if "VETO_" in detailed_order:
            # 移除内部 veto 代码
            detailed_order = _re.sub(r'VETO_[^\s]*', '', detailed_order)

        label_cn_map = {"签": "✅ 建议签署", "改": "🔧 修改后签署", "拖": "⏸️ 暂缓签署", "退": "🚫 建议放弃", "VETO": "🛑 不建议签署"}
        label_cn = label_cn_map.get(decision_label, decision_label)

        report += f"""
## 六、决策建议

### 最终建议：{label_cn}

"""
        # 决策理由
        if decision_reasons:
            for i, reason in enumerate(decision_reasons[:3], 1):
                reason = _clean_internal_names(reason)
                reason = reason.strip().lstrip("、。，")
                if reason:
                    report += f"{i}. {reason}\n"
            report += "\n"

        # 详细决策说明
        if detailed_order:
            # 清理所有内部代号
            detailed_order = _clean_internal_names(detailed_order)
            for kw in ["VETO_戒律1_资本安全垫击穿", "VETO_戒律2_显失公平", "VETO_戒律3_信息不对称",
                       "VETO_戒律4_运营失控", "VETO_戒律5_尾部风险不可逆", "VETO_戒律6_资金失控",
                       "HANDOFF_TO_HUMAN", "六位评审员", "私董会"]:
                detailed_order = detailed_order.replace(kw, "")
            detailed_order = _re.sub(r'\*\*依据\*\*:\s*\n?', '', detailed_order)
            report += f"{detailed_order.strip()}\n\n"

        # 免责声明
        report += f"""
## 七、免责声明

> **本报告为霖信莯信息咨询基于 Qiaoxi Contract-Analyzer 系统生成的商业决策辅助分析，不构成正式的法律意见、财务建议或投资建议。**
>
> 报告中的法律风险分析基于系统内置的中国法律法规数据库，该数据库力求准确但可能存在更新延迟或个别条文解读偏差。报告中的商业分析和推演基于系统对合同文本的自动化解析，不可避免地存在信息不完整的局限。
>
> **您在使用本报告做出商业决策前，建议：**
> - 咨询持有执业资质的律师对合同进行正式法律审查
> - 对交易对手的资质、资产权属等进行独立的尽职调查
> - 结合您自身的商业经验和判断做出最终决策
>
> 霖信莯不对因使用本报告而产生的任何商业后果承担责任。

---

**昆明霖信莯科技有限公司**
*本报告由 Qiaoxi Contract-Analyzer 自动生成*
*报告日期：{datetime.now().strftime('%Y年%m月%d日')}*
"""
        st.session_state.final_report = report
        progress_bar.progress(100, text="完成！")
        status_text.success("审查报告已生成")
        audit.log_state_transition("2_legal_review", "7_report_standard")

        st.session_state.phase = "results"
        st.rerun()

    except Exception as e:
        st.error("处理阶段出错: %s" % str(e))
        st.stop()


# ╔══════════════════════════════════════════════════════════════╗
# ║  STEP 5: 结果展示 + 服务结束（含 15 秒删除倒计时）           ║
# ╚══════════════════════════════════════════════════════════════╝
elif _phase == "results":

    # ── 体验码使用记录 ──
    if st.session_state.get("trial_code") and not st.session_state.get("trial_recorded"):
        try:
            resp = _requests.post(
                "http://localhost:3000/api/trial/use",
                json={"code": st.session_state.trial_code, "product": "qiaoxi"},
                timeout=5
            )
            data = resp.json()
            if data.get("success"):
                st.session_state.trial_recorded = True
                remaining = data.get("remaining", {})
                st.success(f"✅ 体验已记录！乔曦剩余 {remaining.get('qiaoxi', 0)} 次 | 峤远剩余 {remaining.get('qiaoyuan', 0)} 次 | 程晓融剩余 {remaining.get('cxr', 0)} 次")
            else:
                st.warning(f"体验记录失败：{data.get('error', '未知错误')}")
        except Exception:
            st.warning("体验记录失败，请手动联系客服。")

    st.markdown("### 📊 审查报告")

    risks = st.session_state.legal_review.get("risks", []) if st.session_state.legal_review else []
    high_c = len([r for r in risks if r.get("risk_level") == "high"])
    mid_c = len([r for r in risks if r.get("risk_level") == "medium"])
    low_c = len([r for r in risks if r.get("risk_level") == "low"])

    c1, c2, c3, c4 = st.columns(4)
    # 用原始脱敏文本重新精确计数，与报告同步
    _sanitized = st.session_state.contract_sanitized or ""
    _real_count = len(_re.findall(r'第[一二三四五六七八九十百千\d]+条', _sanitized))
    if _real_count == 0:
        _real_count = st.session_state.clause_tree.get("total_clauses", 0) if st.session_state.clause_tree else 0
    c1.metric("条款总数", _real_count)
    c2.metric("🔴 高风险", high_c)
    c3.metric("🟡 中风险", mid_c)
    c4.metric("🟢 低风险", low_c)

    st.divider()
    if st.session_state.final_report:
        st.markdown(st.session_state.final_report)
        st.divider()
        col_md, col_docx = st.columns(2)
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        col_md.download_button(
            "📥 下载报告 (Markdown)",
            data=st.session_state.final_report,
            file_name=f"Qiaoxi_Report_{ts}.md",
            mime="text/markdown",
            use_container_width=True,
        )
        try:
            docx_bytes = _markdown_to_docx(st.session_state.final_report)
            col_docx.download_button(
                "📄 下载报告 (Word)",
                data=docx_bytes,
                file_name=f"Qiaoxi_Report_{ts}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as _e:
            col_docx.warning(f"Word 格式生成失败: {_e}")

    st.divider()

    # ─── 高级服务：Qiaoxi 合同重构 ───
    st.markdown("### 🔧 高级服务：Qiaoxi 合同重构")

    # 付费槽位
    with st.expander("📋 服务说明与收费标准", expanded=False):
        st.markdown("""
**Qiaoxi 合同重构服务**是基于本次审查报告，由 AI 引擎结合客户画像，重新起草一套完整的合同解决方案：

| 交付物 | 说明 |
|--------|------|
| 🔵 **新合同草案** | 针对审查报告所有风险点逐条重构，含重构说明注释 |
| 🟢 **框架协议（代尽调意向书）** | 防御性框架协议，保障签署前的尽调权和零付款原则 |
| 🟡 **尽职调查清单** | 分六类的详细文件核查清单，覆盖法律/财务/资产/政策 |

**收费标准：** ¥30 / 次（含三份文件的生成与一次修改）

如需购买服务，请联系霖信莯咨询获取授权码：
📞 **联系人：余磊 13987671259** ｜ 📧 **邮箱：425448719@qq.com**
""")

    if not st.session_state.reconstruct_paid:
        st.info("⚠️ 此为付费功能。如已获得授权码，请输入后解锁。")
        col_code, col_btn = st.columns([3, 1])
        auth_input = col_code.text_input(
            "输入授权码", key="auth_code_input", placeholder="QIAOXI-XXXX-XXXX 或 CXL-XXXX-XXXX",
            label_visibility="collapsed"
        )
        if col_btn.button("验证", use_container_width=True):
            code_input = auth_input.strip().upper()

            # ── 先验证网站体验码 ──
            if code_input.startswith("CXL-"):
                try:
                    resp = _requests.post(
                        "http://localhost:3000/api/trial/verify",
                        json={"code": code_input},
                        timeout=5
                    )
                    data = resp.json()
                    if data.get("valid") and data.get("remaining", {}).get("qiaoxi", 0) > 0:
                        st.session_state.reconstruct_paid = True
                        st.session_state.trial_code = code_input
                        st.rerun()
                    elif data.get("valid"):
                        st.error("该体验码的乔曦额度已用完，请使用付费授权码。")
                    else:
                        st.error(data.get("error", "体验码无效。"))
                except Exception:
                    st.error("体验码验证失败，请检查网络连接或使用付费授权码。")
            # ── 原有授权码验证 ──
            else:
                valid_codes = {"QIAOXI-RECON-2024", "QIAOXI-DEMO-2024", "QIAOXI-BETA-0001"}
                if code_input in valid_codes:
                    st.session_state.reconstruct_paid = True
                    st.rerun()
                else:
                    st.error("授权码无效，请联系霖信莯咨询获取正式授权码。")

    if st.session_state.reconstruct_paid:
        st.success("✅ 已解锁 Qiaoxi 合同重构服务")
        st.markdown("#### 请选择合同重构方向")
        reconstruct_angle = st.radio(
            "您希望新合同的条款倾向于哪个方向？",
            options=["最大限度维护我方利益", "条款公平偏中性", "兼顾各方利益，便于尽快达成"],
            index=None,
            key="reconstruct_angle_radio",
        )
        if reconstruct_angle:
            st.session_state.reconstruct_angle = reconstruct_angle
            st.markdown(
                f"**已选择：{reconstruct_angle}**\n\n"
                "Qiaoxi 将在精读本次全部审查分析后，针对每一个风险点制定重构方案，"
                "生成新合同草案 + 框架协议 + 尽调清单，共三份 Word 文件。"
            )
            if st.button("🚀 开始合同重构", type="primary", use_container_width=True):
                st.session_state.reconstruct_requested = True
                st.session_state.phase = "reconstructing"
                st.rerun()

    st.divider()
    st.markdown("### ⏹️ 结束服务")

    if not st.session_state.delete_requested:
        st.warning("点击下方按钮后，系统将立即从服务器上物理删除您上传的合同文件及所有中间态数据。此操作不可撤销。")
        if st.button("⏹️ 结束服务并删除我的数据", type="primary", use_container_width=True):
            # 立即删除文件
            import glob
            for fp in st.session_state.session_files:
                try:
                    if os.path.exists(fp):
                        with open(fp, 'wb') as f:
                            f.write(os.urandom(min(os.path.getsize(fp), 4096)))
                        os.unlink(fp)
                except Exception:
                    pass
            for pattern in ["data/*.json"]:
                for f in glob.glob(pattern):
                    try:
                        if os.path.exists(f):
                            os.unlink(f)
                    except Exception:
                        pass
            audit.log_file_deletion(st.session_state.uploaded_file.name if st.session_state.uploaded_file else "unknown", "user_request")
            st.session_state.delete_requested = True
            st.rerun()
    else:
        st.success("✅ 您的合同文件和中间态数据已从服务器上删除。感谢使用霖信莯 Qiaoxi。")
        st.info("如需重新使用本系统，请手动刷新页面。")

    st.divider()
    if st.button("🔄 开始新的审查", type="secondary", use_container_width=True):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.session_state.phase = "consent"
        st.rerun()
        st.rerun()

# ═══════════════════════════════════════════════════════════════
# Phase: reconstructing — 合同重构生成中
# ═══════════════════════════════════════════════════════════════
elif _phase == "reconstructing":
    st.markdown("### 🔧 Qiaoxi 合同重构引擎运行中")
    st.markdown(f"重构方向：**{st.session_state.reconstruct_angle}**")

    prog = st.progress(0, text="准备中…")
    status_text = st.empty()

    def _recon_progress(step: int, total: int, label: str):
        pct = int(step / max(total, 1) * 100) if total > 0 else 0
        prog.progress(pct, text=label)
        status_text.info(f"⏳ {label}")

    # 页面最下方提示框 —— 放在 try 块之前，确保在处理过程中就渲染
    st.markdown("""
<div style="background:#e8f4fd;border:2px solid #1890ff;border-radius:10px;padding:20px 24px;margin:16px 0;text-align:center;">
  <div style="font-size:22px;margin-bottom:8px;">⏳</div>
  <div style="font-size:16px;font-weight:700;color:#0f3460;margin-bottom:6px;">
    Qiaoxi Contract-Analyzer 正在处理您的需求
  </div>
  <div style="font-size:14px;color:#555;line-height:1.8;">
    系统正在为您重构合同条款及附件，请耐心等待，预计需要 60–120 秒。<br/>
    <strong style="color:#cf1322;">请不要刷新页面，谢谢您的耐心。</strong><br/>
    您也可以回到本页面上方查看进度条。
  </div>
</div>
""", unsafe_allow_html=True)

    try:
        engine = ReconstructionEngine()
        outputs = engine.run(
            final_report=st.session_state.final_report or "",
            clause_tree=st.session_state.clause_tree or {},
            legal_review=st.session_state.legal_review or {},
            audits=st.session_state.six_audits or [],
            decision=st.session_state.decision or {},
            profile=st.session_state.profile or {},
            reconstruct_angle=st.session_state.reconstruct_angle,
            progress_callback=_recon_progress,
        )
        st.session_state.reconstruct_outputs = outputs
        prog.progress(100, text="重构完成！")

        # 动态成功提示
        files_gen = outputs.get("files_generated", {})
        gen_count = sum(1 for v in files_gen.values() if v)
        fw = files_gen.get("framework_agreement")
        dd = files_gen.get("dd_checklist")
        item_list = ["✅ 新合同草案"]
        if fw:
            item_list.append("✅ 框架协议")
        if dd:
            item_list.append("✅ 尽调清单")
        status_text.success(f"重构完成 — 共生成 {gen_count} 份文件：" + " · ".join(item_list))

        audit.log_state_transition("7_report_standard", "8_report_advanced")
        st.session_state.phase = "reconstruct_done"
        st.rerun()

    except Exception as _exc:
        prog.progress(0, text="")
        st.error(f"重构引擎运行失败：{_exc}")
        st.code(traceback.format_exc())
        if st.button("返回报告页面", type="secondary"):
            st.session_state.phase = "results"
            st.rerun()

# ═══════════════════════════════════════════════════════════════
# Phase: reconstruct_done — 下载重构结果
# ═══════════════════════════════════════════════════════════════
elif _phase == "reconstruct_done":
    outputs = st.session_state.reconstruct_outputs or {}
    angle = st.session_state.reconstruct_angle or "未选择"
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')

    files_gen = outputs.get("files_generated", {})
    gen_fw = files_gen.get("framework_agreement")
    gen_dd = files_gen.get("dd_checklist")
    gen_count = 1 + (1 if gen_fw else 0) + (1 if gen_dd else 0)

    st.markdown("### ✅ 合同重构完成")
    item_desc = "新合同草案"
    if gen_fw:
        item_desc += " + 框架协议"
    if gen_dd:
        item_desc += " + 尽调清单"
    st.success(f"已按「{angle}」方向完成重构，共生成 {gen_count} 份文件：{item_desc}。")
    if not gen_fw and not gen_dd:
        st.caption("DeepSeek 综合判断本合同为简单合同，不需要框架协议和尽调清单。仅生成新合同草案。")
    elif not gen_fw:
        st.caption("DeepSeek 综合判断本合同的复杂程度不需要框架协议（尽调清单已根据实际需要生成）。")
    elif not gen_dd:
        st.caption("DeepSeek 综合判断本合同的复杂程度不需要单独的尽调清单。")

    st.markdown("下方文件均为可下载的 Word 格式。")
    st.divider()

    def _recon_docx(md_text: str, doc_title: str) -> bytes:
        header = f"# {doc_title}\n\n*生成方向：{angle}*\n*生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}*\n*霖信莯咨询 · Qiaoxi Contract-Analyzer*\n\n---\n\n"
        return _markdown_to_docx(header + (md_text or "（内容未生成）"))

    # 动态三列，跳过不需要的文件
    col_slots = [1]  # 新合同草案始终占一列
    if gen_fw:
        col_slots.append(2)
    if gen_dd:
        col_slots.append(3)
    cols = st.columns(len(col_slots))
    col_idx = 0

    # ── 文件1：新合同草案（始终）──
    with cols[col_idx]:
        st.markdown("#### 📄 新合同草案")
        st.markdown("针对审查报告所有风险点逐条重构，含【重构说明】注释。")
        contract_draft = outputs.get("contract_draft", "")
        if contract_draft:
            try:
                docx_bytes = _recon_docx(contract_draft, "Qiaoxi 重构合同草案")
                st.download_button(
                    "📥 下载新合同草案 (.docx)",
                    data=docx_bytes,
                    file_name=f"Qiaoxi_重构合同_{ts}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary",
                )
            except Exception as _e:
                st.warning(f"Word 转换失败，提供 Markdown 版：{_e}")
                st.download_button(
                    "📥 下载新合同草案 (.md)",
                    data=contract_draft,
                    file_name=f"Qiaoxi_重构合同_{ts}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )
        else:
            st.warning("新合同草案未生成")
    col_idx += 1

    # ── 文件2：框架协议（条件）──
    if gen_fw:
        with cols[col_idx]:
            st.markdown("#### 📋 框架协议")
            st.markdown("防御性框架协议（代尽调意向书），保障甲方付款前的全部权利。")
            framework = outputs.get("framework_agreement", "")
            if framework:
                try:
                    docx_bytes = _recon_docx(framework, "Qiaoxi 框架协议（代尽调意向书）")
                    st.download_button(
                        "📥 下载框架协议 (.docx)",
                        data=docx_bytes,
                        file_name=f"Qiaoxi_框架协议_{ts}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="primary",
                    )
                except Exception as _e:
                    st.warning(f"Word 转换失败，提供 Markdown 版：{_e}")
                    st.download_button(
                        "📥 下载框架协议 (.md)",
                        data=framework,
                        file_name=f"Qiaoxi_框架协议_{ts}.md",
                        mime="text/markdown",
                        use_container_width=True,
                    )
            else:
                st.warning("框架协议未生成")
        col_idx += 1

    # ── 文件3：尽调清单（条件）──
    if gen_dd:
        with cols[col_idx]:
            st.markdown("#### 📑 尽职调查清单")
            st.markdown("分六类的详细文件核查清单，含优先级标注。")
            dd_checklist = outputs.get("dd_checklist", "")
            if dd_checklist:
                try:
                    docx_bytes = _recon_docx(dd_checklist, "Qiaoxi 尽职调查清单（建议）")
                    st.download_button(
                        "📥 下载尽调清单 (.docx)",
                        data=docx_bytes,
                        file_name=f"Qiaoxi_尽调清单_{ts}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="primary",
                    )
                except Exception as _e:
                    st.warning(f"Word 转换失败，提供 Markdown 版：{_e}")
                    st.download_button(
                        "📥 下载尽调清单 (.md)",
                        data=dd_checklist,
                        file_name=f"Qiaoxi_尽调清单_{ts}.md",
                        mime="text/markdown",
                        use_container_width=True,
                    )
            else:
                st.warning("尽调清单未生成")

    st.divider()
    col_back, col_new = st.columns(2)
    if col_back.button("← 返回分析报告", use_container_width=True):
        st.session_state.phase = "results"
        st.rerun()
    if col_new.button("🔄 开始新的审查", type="secondary", use_container_width=True):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()
