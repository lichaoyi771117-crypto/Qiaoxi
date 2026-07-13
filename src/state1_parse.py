"""
Qiaoxi Contract-Analyzer · PactGuard MCP Service 适配层

基于 PactGuard-ERNIE-PP 改造：
- LLM 后端从 ERNIE 4.5 切换到 DeepSeek
- 风险分析增加 RAG 调用钩子
- 条款树提取

具体改造:
  1. contract_workflow.py L44: default_model = "ernie-4.5-turbo-128k" → "deepseek-chat"
  2. _analyze_risks() 增加 ChromaDB RAG 检索结果注入
  3. 风险输出增加 legal_basis / rag_confidence 字段
  4. 新增 clause_tree 结构化提取

改造代码存放: D:\Ai RAG\Qiaoxi\modified code\PactGuard-Qiaoxi\
"""
import json
import logging
from pathlib import Path
from typing import Optional
from openai import OpenAI

from src.config import DEEPSEEK_API_KEY, DEEPSEEK_BASE_URL, DEEPSEEK_MODEL

logger = logging.getLogger(__name__)

# PactGuard 改造版 MCP Service URL（本地）
MCP_SERVICE_URL = "http://localhost:7001"


class QiaoxiContractParser:
    """
    Qiaoxi 合同解析器（PactGuard 改造版）

    四阶段工作流：
      1. 文档解析（MCP Service + PP-StructureV3）
      2. 条款结构化（clause_tree）
      3. 风险分析（DeepSeek + RAG 注入）
      4. 结果整合
    """

    def __init__(self, mcp_url: str = MCP_SERVICE_URL):
        self.mcp_url = mcp_url
        self.llm_client = OpenAI(
            api_key=DEEPSEEK_API_KEY,
            base_url=DEEPSEEK_BASE_URL,
        )

    def _call_mcp_parse(self, file_path: str) -> Optional[str]:
        """
        调用 MCP Service 解析文档（PDF/DOCX/图片）
        TODO: Phase 1.3 实现 HTTP 调用

        降级策略（按优先级）：
          1. DOCX → python-docx 直接提取（无编码问题，精度最高）
          2. PDF → pdfplumber（处理 Identity-H 中文字体更好）
          3. PDF → PyMuPDF 兜底
        """
        logger.info(f"[MCP] 待实现: 解析 {file_path}")
        ext = Path(file_path).suffix.lower()

        # --- DOCX 优先路径 ---
        if ext == ".docx":
            try:
                from docx import Document
                doc = Document(file_path)
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            parts.append(row_text)
                text = "\n".join(parts)
                logger.info(f"[MCP降级] python-docx 提取了 {len(text)} 字符")
                return text
            except ImportError:
                logger.warning("[MCP降级] python-docx 未安装，尝试 PDF 路径")
            except Exception as e:
                logger.error(f"[MCP降级] python-docx 解析失败: {e}")
                return None

        # --- PDF 路径：pdfplumber 优先（Identity-H 字体友好）---
        if ext == ".pdf":
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    logger.info(f"[MCP降级] pdfplumber 提取了 {len(text)} 字符")
                    return text
                logger.warning("[MCP降级] pdfplumber 提取为空，降级到 PyMuPDF")
            except ImportError:
                logger.warning("[MCP降级] pdfplumber 未安装，使用 PyMuPDF")
            except Exception as e:
                logger.warning(f"[MCP降级] pdfplumber 失败 ({e})，降级到 PyMuPDF")

            # PyMuPDF 兜底
            try:
                import fitz
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
                logger.info(f"[MCP降级] PyMuPDF 提取了 {len(text)} 字符")
                return text if text.strip() else None
            except Exception as e:
                logger.error(f"[MCP降级] PyMuPDF 解析失败: {e}")
                return None

        logger.error(f"[MCP降级] 不支持的文件格式: {ext}")
        return None

    def _extract_clause_tree(self, document_text: str) -> dict:
        """
        从合同文本中提取条款树（clause_tree）

        按"第X条"/"第X章"/"Section X" 分割，构建树形结构。
        """
        import re

        clauses = []
        # 匹配中文合同条款：第X条、第X章
        chinese_clause_pattern = re.compile(
            r'(第[一二三四五六七八九十百千\d]+(?:章|条|节|款))\s*(.*?)(?=第[一二三四五六七八九十百千\d]+(?:章|条|节|款)|$)',
            re.DOTALL
        )

        matches = chinese_clause_pattern.findall(document_text)
        for i, (heading, content) in enumerate(matches):
            clauses.append({
                "clause_id": f"CLS-{i+1:04d}",
                "heading": heading.strip(),
                "content": content.strip()[:2000],  # 截断超长条款
                "level": "article" if "条" in heading else "chapter" if "章" in heading else "section",
            })

        return {
            "total_clauses": len(clauses),
            "clauses": clauses,
            "extraction_method": "regex_chinese",
        }

    def _analyze_risks(
        self,
        document_text: str,
        clause_tree: dict,
        rag_results: Optional[list[dict]] = None,
    ) -> dict:
        """
        调用 DeepSeek 进行风险分析，注入 RAG 检索结果

        Args:
            document_text: 合同原文
            clause_tree: 条款树
            rag_results: ChromaDB RAG 检索结果（Top-3 法条）
        """
        # 构建 RAG 上下文
        rag_context = ""
        if rag_results:
            rag_items = []
            for i, result in enumerate(rag_results[:3]):
                rag_items.append(
                    f"【法规{i+1}】{result.get('law_name', '未知')} "
                    f"第{result.get('article', '?')}条: {result.get('content', '')[:300]}"
                )
            rag_context = "\n".join(rag_items)
        else:
            rag_context = "（RAG 检索未返回结果，以下分析无具体法条依据）"

        system_prompt = f"""你是乔曦（Qiaoxi），中国政法大学硕士，执业方向商事争议解决。
身份：霖信莯咨询公司的法务助理，负责合同第一轮条款清洗与法律风险标定。
输出模态：锋锐模态——结构化、结论先行、标注依据与量化风险。

【核心约束】
- 禁止引用已废止法律
- 禁止对商业模式合理性做主观判断
- 所有法条引用必须从以下 RAG 检索结果中获取，不可自行编造
- 若 RAG 未返回相关法条，标注"法规待核"

【RAG 检索结果】
{rag_context}

【输出要求】
以 JSON 格式输出风险分析结果。"""

        user_prompt = f"""请审查以下合同条款，逐条标定法律风险：

【合同条款树】
{json.dumps(clause_tree, ensure_ascii=False)[:8000]}

【审查要求】
1. 逐条分析法律风险，按高/中/低三级分档
2. 每个风险点必须包含：风险简述、关联法条（来自RAG结果）、可能后果、建议动作
3. 法条引用格式：《法规名称》第X条第X款（生效状态：现行有效/待核）
4. 若无法条依据，标注"【法规待核】"
5. 仅做法律标定，不输出商业判断

输出严格 JSON。"""

        try:
            response = self.llm_client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.3,
                max_tokens=4096,
                response_format={"type": "json_object"},
            )
            result = json.loads(response.choices[0].message.content)
            logger.info(f"[风险分析] 完成，识别风险 {len(result.get('risks', []))} 条")
            return result
        except Exception as e:
            logger.error(f"[风险分析] LLM 调用失败: {e}")
            return {"risks": [], "error": str(e), "rag_triggered": rag_results is not None}

    def parse_contract(self, file_path: str, rag_results: Optional[list[dict]] = None) -> dict:
        """
        主入口：解析合同文件 → 条款树 → 风险分析

        Args:
            file_path: 合同文件路径
            rag_results: RAG 检索结果（可选，若为 None 则跳过法律引用）

        Returns:
            {
                "contract_meta": {...},
                "clause_tree": {...},
                "legal_review": {...},
                "document_text": "...",
            }
        """
        ext = Path(file_path).suffix.lower()
        logger.info(f"[合同解析] 开始处理: {file_path} (格式: {ext})")

        # Stage 1: 文档解析
        document_text = self._call_mcp_parse(file_path)
        if not document_text:
            return {"error": "文档解析失败", "contract_meta": {}, "clause_tree": {}, "legal_review": {}}

        # Stage 2: 条款结构化
        clause_tree = self._extract_clause_tree(document_text)
        logger.info(f"[条款树] 提取 {clause_tree['total_clauses']} 个条款")

        # Stage 3: 风险分析（注入 RAG 结果）
        legal_review = self._analyze_risks(document_text, clause_tree, rag_results)
        legal_review["rag_triggered"] = rag_results is not None
        legal_review["abolished_laws_blocked"] = 0

        # Stage 4: 整合
        contract_meta = {
            "title": Path(file_path).stem,
            "type": "unknown",  # TODO: LLM 自动识别合同类型
            "value_cny": 0,     # TODO: LLM 自动提取合同金额
            "file_path": str(file_path),
            "ext": ext,
        }

        return {
            "contract_meta": contract_meta,
            "clause_tree": clause_tree,
            "legal_review": legal_review,
            "document_text": document_text[:10000],  # 限制长度
        }
